"""
WORD 增强处理器
提取文本、表格转Markdown、图片OCR，统一输出Markdown格式
"""

import os
from pathlib import Path
from typing import Dict, Any, List
import logging

from docx import Document as DocxDocument

from app.services.ocr_service import OCRService

logger = logging.getLogger(__name__)


def _table_to_markdown(table) -> str:
    """将 Word 表格转换为 Markdown 格式"""
    rows_data = []
    
    for row in table.rows:
        row_data = [cell.text.strip() if cell.text else "" for cell in row.cells]
        rows_data.append(row_data)
    
    if not rows_data:
        return ""
    
    # 第一行作为表头
    header = rows_data[0] if rows_data else []
    header_row = "| " + " | ".join(header) + " |"
    separator = "| " + " | ".join(["---"] * len(header)) + " |"
    
    markdown_rows = [header_row, separator]
    
    # 数据行
    for row_data in rows_data[1:]:
        # 补齐列数
        while len(row_data) < len(header):
            row_data.append("")
        row_md = "| " + " | ".join(row_data[:len(header)]) + " |"
        markdown_rows.append(row_md)
    
    return "\n".join(markdown_rows)


def _extract_images_from_docx(docx_path: str) -> List[Dict[str, Any]]:
    """从 Word 文档中提取所有图片"""
    images_info = []
    
    try:
        # docx 文件实际上是 zip，包含图片在 word/media/ 目录
        import zipfile
        import tempfile
        from pathlib import Path
        
        with zipfile.ZipFile(docx_path, 'r') as zip_ref:
            image_files = [f for f in zip_ref.namelist() if f.startswith('word/media/')]
            
            for img_file in image_files:
                img_name = Path(img_file).name
                # 提取图片到临时文件
                with tempfile.NamedTemporaryFile(delete=False, suffix=Path(img_name).suffix) as tmp:
                    tmp.write(zip_ref.read(img_file))
                    tmp_path = tmp.name
                    
                    images_info.append({
                        "name": img_name,
                        "path": tmp_path,
                        "media_path": img_file
                    })
    except Exception as e:
        logger.warning(f"提取 Word 图片失败: {e}")
    
    return images_info


def process_word_to_markdown(file_path: str) -> Dict[str, Any]:
    """将 Word 文档转换为 Markdown 格式"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Word文件不存在: {file_path}")
    
    doc = DocxDocument(file_path)
    ocr_service = OCRService()
    
    markdown_parts: List[str] = []
    images_processed = []
    
    # 提取图片信息
    images_info = _extract_images_from_docx(file_path)
    
    # 遍历文档中的段落和表格
    # 使用更简单的方法：按顺序处理段落和表格
    para_idx = 0
    table_idx = 0
    
    # 获取所有段落的样式信息
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 检查段落样式
            style_name = para.style.name if para.style else ""
            if 'Heading' in style_name or 'heading' in style_name.lower():
                # 标题样式
                if 'Heading 1' in style_name or 'Heading1' in style_name:
                    markdown_parts.append(f"# {text}\n")
                elif 'Heading 2' in style_name or 'Heading2' in style_name:
                    markdown_parts.append(f"## {text}\n")
                elif 'Heading 3' in style_name or 'Heading3' in style_name:
                    markdown_parts.append(f"### {text}\n")
                elif 'Heading 4' in style_name or 'Heading4' in style_name:
                    markdown_parts.append(f"#### {text}\n")
                else:
                    markdown_parts.append(f"## {text}\n")  # 默认二级标题
            else:
                markdown_parts.append(f"{text}\n")
    
    # 处理表格
    for table in doc.tables:
        table_md = _table_to_markdown(table)
        if table_md:
            markdown_parts.append(f"\n### 表格 {table_idx + 1}\n\n{table_md}\n\n")
            table_idx += 1
    
    # 处理图片 OCR
    if images_info:
        markdown_parts.append("\n### 图片OCR内容\n\n")
        for img_info in images_info:
            try:
                ocr_result = ocr_service.extract_text_from_image(img_info["path"])
                if isinstance(ocr_result, dict):
                    ocr_text = ocr_result.get("text", "").strip()
                else:
                    ocr_text = str(ocr_result).strip()
                
                if ocr_text:
                    markdown_parts.append(f"**图片 {img_info['name']}**:\n{ocr_text}\n\n")
                    images_processed.append({
                        "name": img_info["name"],
                        "ocr_text": ocr_text
                    })
                
                # 清理临时文件
                try:
                    os.unlink(img_info["path"])
                except Exception:
                    pass
            except Exception as e:
                logger.warning(f"图片 OCR 失败 {img_info['name']}: {e}")
                try:
                    os.unlink(img_info["path"])
                except Exception:
                    pass
    
    final_markdown = "\n".join(markdown_parts)
    
    return {
        "file_type": "word",
        "content": final_markdown,
        "metadata": {
            "paragraph_count": len([p for p in doc.paragraphs if p.text.strip()]),
            "table_count": len(doc.tables),
            "image_count": len(images_processed),
            "file_size": os.path.getsize(file_path)
        },
        "tables": [{"index": i, "markdown": _table_to_markdown(t)} for i, t in enumerate(doc.tables)],
        "images_processed": images_processed
    }

